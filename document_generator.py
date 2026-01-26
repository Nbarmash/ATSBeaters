
import os
from docx import Document
from docx.shared import Inches, Pt
from config import OUTPUT_DIR, logger

class DocumentGenerator:
    def markdown_to_docx(self, markdown_content: str, filename: str):
        """
        Converts markdown-style text into an ATS-optimized Word document.
        Standardizes fonts and margins for maximum parsability.
        """
        logger.info(f"Generating DOCX: {filename}")
        
        doc = Document()
        
        # Define base styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Set standard 1-inch margins (Industry standard for ATS)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Basic markdown parsing logic (Simulated for this script)
        lines = markdown_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        output_path = os.path.join(OUTPUT_DIR, filename)
        doc.save(output_path)
        logger.info(f"Document saved to {output_path}")
        return output_path
